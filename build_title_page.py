"""Generate the AAA title page docx for the ARS-VG manuscript.

Structure follows the AAA-provided template exactly:
    [Title, 12pt Times New Roman]
    [blank]
    [Author / Affiliation blocks, one per author]
    Running Head: [...]
    [Acknowledgments]
    [Financial COI disclosure]
    [Other disclosures, if applicable]
    [Full author block: NAME, AFFILIATION, COLLEGE, DEPT, CITY, ST, COUNTRY; ...]
    Data Availability: [...]
    JEL Classifications: [...]
    Keywords: [...]
    Does this article have supplemental material(s)? [Yes/No]
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE = (
    "ARS-VG: A Graph-Theoretic Intelligent Artifact for Joint Detection of "
    "Accrual-Based and Real-Activities Earnings Management Substitution in "
    "SEC-Registered Firms — A Design Science Research Approach"
)

RUNNING_HEAD = (
    "ARS-VG: Graph-Based Detection of AEM–REM Substitution in "
    "SEC-Registered Firms"
)
assert len(RUNNING_HEAD) <= 115, f"Running head too long: {len(RUNNING_HEAD)}"

doc = Document()

core = doc.core_properties
core.author = "apoorv"
core.last_modified_by = "apoorv"
core.title = TITLE

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
fmt = style.paragraph_format
fmt.space_before = Pt(0)
fmt.space_after = Pt(0)
fmt.line_spacing = 2.0

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_blank():
    doc.add_paragraph("")


def add_par(text=None, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, runs=None):
    """Add a single-run paragraph or, if `runs` is given, a multi-run paragraph.

    `runs` is a list of (text, bold) tuples.
    """
    p = doc.add_paragraph()
    p.alignment = align
    if runs is None:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    else:
        for rtext, rbold in runs:
            r = p.add_run(rtext)
            r.bold = rbold
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
    return p


# --- Title (12pt Times New Roman, centered) ---------------------------------
add_par(TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()
add_blank()

# --- Author / Affiliation blocks --------------------------------------------
add_par("Apoorv Saxena", align=WD_ALIGN_PARAGRAPH.CENTER)
add_par(
    "College of Management, Yuan Ze University",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_blank()

add_par("Yan-Jie Yang", align=WD_ALIGN_PARAGRAPH.CENTER)
add_par(
    "College of Management, Yuan Ze University",
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_blank()
add_blank()

# --- Running head ------------------------------------------------------------
add_par(
    runs=[
        ("Running Head: ", False),
        (RUNNING_HEAD, False),
    ]
)
add_blank()

# --- Acknowledgments --------------------------------------------------------
add_par(
    "The authors thank colleagues for helpful comments on earlier versions of "
    "this paper. Any remaining errors are our own. This research received no "
    "specific grant from any funding agency in the public, commercial, or "
    "not-for-profit sectors."
)
add_blank()

# --- Financial COI Disclosure -----------------------------------------------
add_par(
    "Apoorv Saxena and Yan-Jie Yang have no financial conflicts of interest "
    "related to this research."
)
add_blank()

# --- Full author / affiliation block ----------------------------------------
add_par(
    "Apoorv Saxena, Yuan Ze University, College of Management, Chung-Li, "
    "Taoyuan, Taiwan (ORCID: 0009-0002-4268-2436); Yan-Jie Yang "
    "(corresponding author; yanjie@saturn.yzu.edu.tw), Yuan Ze University, "
    "College of Management, Chung-Li, Taoyuan, Taiwan "
    "(ORCID: 0000-0001-7346-5050)."
)
add_blank()

# --- Data Availability ------------------------------------------------------
add_par(
    runs=[
        ("Data Availability: ", False),
        (
            "Data and code supporting the findings of this study are "
            "available from the authors on reasonable request.",
            False,
        ),
    ]
)
add_blank()

# --- JEL Classifications ----------------------------------------------------
add_par(
    runs=[
        ("JEL Classifications: ", False),
        ("M41; M42; C45; G30.", False),
    ]
)
add_blank()

# --- Keywords ---------------------------------------------------------------
add_par(
    runs=[
        ("Keywords: ", False),
        (
            "Earnings management; AEM–REM substitution; Agency theory; "
            "Design Science Research; Financial knowledge graph; Graph-based "
            "anomaly detection; Forensic accounting.",
            False,
        ),
    ]
)
add_blank()

# --- Supplemental material --------------------------------------------------
add_par(
    "Does this article have supplemental material(s) that are intended for "
    "publication? No."
)

out_path = "/home/user/cli-2/title_page.docx"
doc.save(out_path)
print(f"Wrote {out_path}")
print(f"Running head length: {len(RUNNING_HEAD)} chars")
